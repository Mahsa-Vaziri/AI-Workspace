import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentObject
from docx.text.paragraph import Paragraph


def save_word_document(file_name: str, content: str) -> Path:
    """Convert Markdown content into a formatted Word document."""

    output_folder = Path("output")
    output_folder.mkdir(exist_ok=True)

    file_path = output_folder / file_name

    document = Document()

    document.add_heading(
        "Advisor AI Studio Report",
        level=0,
    )

    lines = content.splitlines()
    line_index = 0

    while line_index < len(lines):
        line = lines[line_index].strip()

        # رد شدن از خط‌های خالی
        if not line:
            line_index += 1
            continue

        # Heading سطح 3
        if line.startswith("### "):
            heading_text = line.removeprefix("### ").strip()

            document.add_heading(
                heading_text,
                level=3,
            )

        # Heading سطح 2
        elif line.startswith("## "):
            heading_text = line.removeprefix("## ").strip()

            document.add_heading(
                heading_text,
                level=2,
            )

        # Heading سطح 1
        elif line.startswith("# "):
            heading_text = line.removeprefix("# ").strip()

            document.add_heading(
                heading_text,
                level=1,
            )

        # تشخیص جدول Markdown
        elif line.startswith("|"):
            table_lines = []

            while (
                line_index < len(lines)
                and lines[line_index].strip().startswith("|")
            ):
                table_lines.append(
                    lines[line_index].strip()
                )

                line_index += 1

            add_markdown_table(
                document,
                table_lines,
            )

            continue

        # تشخیص Bullet List
        elif re.match(r"^[-*+]\s+", line):
            bullet_text = re.sub(
                r"^[-*+]\s+",
                "",
                line,
            )

            paragraph = document.add_paragraph(
                style="List Bullet",
            )

            add_formatted_text(
                paragraph,
                bullet_text,
            )

        # تشخیص Numbered List
        elif re.match(r"^\d+\.\s+", line):
            numbered_text = re.sub(
                r"^\d+\.\s+",
                "",
                line,
            )

            paragraph = document.add_paragraph(
                style="List Number",
            )

            add_formatted_text(
                paragraph,
                numbered_text,
            )

        # پاراگراف عادی
        else:
            paragraph = document.add_paragraph()

            add_formatted_text(
                paragraph,
                line,
            )

        line_index += 1

    document.save(file_path)

    return file_path


def add_formatted_text(
    paragraph: Paragraph,
    text: str,
) -> None:
    """Convert Markdown bold and italic text into Word formatting."""

    markdown_pattern = r"(\*\*.+?\*\*|\*.+?\*)"

    text_parts = re.split(
        markdown_pattern,
        text,
    )

    for text_part in text_parts:
        if not text_part:
            continue

        # Bold: **example**
        if (
            text_part.startswith("**")
            and text_part.endswith("**")
        ):
            clean_text = text_part[2:-2]

            text_run = paragraph.add_run(clean_text)
            text_run.bold = True

        # Italic: *example*
        elif (
            text_part.startswith("*")
            and text_part.endswith("*")
        ):
            clean_text = text_part[1:-1]

            text_run = paragraph.add_run(clean_text)
            text_run.italic = True

        # متن عادی
        else:
            paragraph.add_run(text_part)


def add_markdown_table(
    document: DocumentObject,
    table_lines: list[str],
) -> None:
    """Convert Markdown table lines into a Word table."""

    valid_rows: list[list[str]] = []

    for line in table_lines:
        cells = [
            cell.strip()
            for cell in line.strip("|").split("|")
        ]

        if is_separator_row(cells):
            continue

        valid_rows.append(cells)

    if not valid_rows:
        return

    column_count = len(valid_rows[0])

    table = document.add_table(
        rows=1,
        cols=column_count,
    )

    table.style = "Table Grid"

    header_cells = table.rows[0].cells

    # ساخت Header جدول
    for column_index, value in enumerate(valid_rows[0]):
        add_formatted_text(
            header_cells[column_index].paragraphs[0],
            value,
        )

    # ساخت ردیف‌های جدول
    for row_values in valid_rows[1:]:
        row_cells = table.add_row().cells

        for column_index, value in enumerate(row_values):
            if column_index < column_count:
                add_formatted_text(
                    row_cells[column_index].paragraphs[0],
                    value,
                )


def is_separator_row(cells: list[str]) -> bool:
    """Check whether a Markdown row is a table separator."""

    if not cells:
        return False

    for cell in cells:
        cleaned_cell = (
            cell
            .replace(":", "")
            .replace("-", "")
            .strip()
        )

        if cleaned_cell:
            return False

    return True